import os
import sys
import types
import unittest
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import Mock, patch

from fastapi.testclient import TestClient


TEST_JWT_SECRET = "test-only-jwt-secret-key-with-at-least-sixty-four-characters-123456"
os.environ.setdefault("JWT_SECRET_KEY", TEST_JWT_SECRET)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
BACKEND_DIR = PROJECT_ROOT / "backend"
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

import main
from auth import get_current_user


class ApiTestCase(unittest.TestCase):
    def setUp(self):
        self.cleanup_patcher = patch.object(main, "cleanup_retained_data", return_value=None)
        self.rate_limit_patcher = patch.object(main.rate_limiter, "enforce", return_value=None)
        self.cleanup_patcher.start()
        self.rate_limit_patcher.start()
        main.app.dependency_overrides.clear()
        self.client = TestClient(main.app)
        self.client.__enter__()

    def tearDown(self):
        self.client.__exit__(None, None, None)
        main.app.dependency_overrides.clear()
        self.rate_limit_patcher.stop()
        self.cleanup_patcher.stop()

    @staticmethod
    def user(user_id=7, username="demo_user"):
        return SimpleNamespace(id=user_id, username=username, is_active=True)

    def authenticate(self, user=None):
        current_user = user or self.user()
        main.app.dependency_overrides[get_current_user] = lambda: current_user
        return current_user

    def test_register_returns_token_and_cookie(self):
        user = self.user()
        with (
            patch.dict(os.environ, {"COOKIE_SECURE": "false"}),
            patch.object(main, "create_user", return_value=user) as create_user,
            patch.object(main, "create_access_token", return_value="register-token"),
        ):
            response = self.client.post(
                "/api/auth/register",
                json={"username": "  demo_user  ", "password": "password123"},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["access_token"], "register-token")
        self.assertEqual(response.cookies.get("access_token"), "register-token")
        self.assertIn("HttpOnly", response.headers["set-cookie"])
        self.assertNotIn("; Secure", response.headers["set-cookie"])
        create_user.assert_called_once_with("demo_user", "password123")

    def test_secure_cookie_can_be_enabled(self):
        user = self.user()
        with (
            patch.dict(os.environ, {"COOKIE_SECURE": "true"}),
            patch.object(main, "create_user", return_value=user),
            patch.object(main, "create_access_token", return_value="secure-token"),
        ):
            response = self.client.post(
                "/api/auth/register",
                json={"username": "secure_user", "password": "password123"},
            )

        self.assertEqual(response.status_code, 200)
        self.assertIn("; Secure", response.headers["set-cookie"])

    def test_register_duplicate_username_returns_409(self):
        with patch.object(main, "create_user", side_effect=ValueError("exists")):
            response = self.client.post(
                "/api/auth/register",
                json={"username": "duplicate", "password": "password123"},
            )

        self.assertEqual(response.status_code, 409)
        self.assertEqual(response.json()["detail"], "username already exists")

    def test_login_success_returns_token(self):
        user = self.user()
        with (
            patch.object(main, "authenticate", return_value=user) as authenticate,
            patch.object(main, "create_access_token", return_value="login-token"),
        ):
            response = self.client.post(
                "/api/auth/login",
                json={"username": "  demo_user ", "password": "password123"},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.cookies.get("access_token"), "login-token")
        authenticate.assert_called_once_with("demo_user", "password123")

    def test_login_invalid_password_returns_401(self):
        with patch.object(main, "authenticate", return_value=None):
            response = self.client.post(
                "/api/auth/login",
                json={"username": "demo_user", "password": "wrongpass"},
            )

        self.assertEqual(response.status_code, 401)
        self.assertEqual(response.json()["detail"], "invalid username or password")

    def test_logout_clears_cookie(self):
        self.client.cookies.set("access_token", "old-token")
        response = self.client.post("/api/auth/logout")

        self.assertEqual(response.status_code, 200)
        self.assertIn("access_token=", response.headers["set-cookie"])
        self.assertIn("Max-Age=0", response.headers["set-cookie"])

    def test_protected_endpoint_requires_authentication(self):
        response = self.client.get("/api/auth/me")

        self.assertEqual(response.status_code, 401)
        self.assertEqual(response.json()["detail"], "authentication required")

    def test_current_user_endpoint(self):
        user = self.authenticate()
        response = self.client.get("/api/auth/me")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"], {"id": user.id, "username": user.username})

    def test_resume_text_upload(self):
        self.authenticate()
        response = self.client.post(
            "/api/resume/upload",
            files={"file": ("resume.txt", "Python backend experience", "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["content"], "Python backend experience")

    def test_resume_export_text_preserves_extension(self):
        self.authenticate()
        response = self.client.post(
            "/api/resume/export",
            data={"optimized_text": "优化后的简历内容"},
            files={"file": ("resume.txt", "原始简历".encode("utf-8"), "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.headers["content-type"].startswith("text/plain"))
        self.assertIn(".txt", response.headers["content-disposition"])
        self.assertEqual(response.content.decode("utf-8-sig"), "优化后的简历内容")
        self.assertTrue(response.content.startswith(b"\xef\xbb\xbf"))

    def test_resume_export_docx_keeps_docx_readable(self):
        from docx import Document

        source = Document()
        source.sections[0].top_margin = 123456
        source.add_heading("原始简历", level=1)
        source.add_paragraph("原始经历")
        source_bytes = BytesIO()
        source.save(source_bytes)

        self.authenticate()
        response = self.client.post(
            "/api/resume/export",
            data={"optimized_text": "优化后的简历\n新的经历"},
            files={
                "file": (
                    "resume.docx",
                    source_bytes.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn(".docx", response.headers["content-disposition"])
        exported = Document(BytesIO(response.content))
        self.assertIn("优化后的简历", "\n".join(paragraph.text for paragraph in exported.paragraphs))
        self.assertEqual(exported.sections[0].top_margin, source.sections[0].top_margin)

    def test_resume_export_pdf_preserves_pdf_type(self):
        import PyPDF2

        self.authenticate()
        source_pdf = main._build_pdf("原始简历", 595.28, 841.89)
        response = self.client.post(
            "/api/resume/export",
            data={"optimized_text": "优化后的 PDF 简历"},
            files={"file": ("resume.pdf", source_pdf, "application/pdf")},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["content-type"], "application/pdf")
        self.assertTrue(response.content.startswith(b"%PDF-"))
        exported_pdf = PyPDF2.PdfReader(BytesIO(response.content))
        self.assertGreater(len(exported_pdf.pages), 0)
        self.assertGreater(len(response.content), 5_000)
        self.assertTrue(b"/FontFile2" in response.content or b"/FontFile3" in response.content)

    def test_resume_export_rejects_unsupported_extension(self):
        self.authenticate()
        response = self.client.post(
            "/api/resume/export",
            data={"optimized_text": "optimized"},
            files={"file": ("resume.exe", b"not a resume", "application/octet-stream")},
        )

        self.assertEqual(response.status_code, 400)

    def test_resume_export_requires_authentication(self):
        response = self.client.post(
            "/api/resume/export",
            data={"optimized_text": "optimized"},
        )

        self.assertEqual(response.status_code, 401)

    def test_resume_optimization_uses_cleaned_input(self):
        self.authenticate()
        with (
            patch.object(main, "enforce_llm_limits", return_value=None),
            patch.object(main, "optimize_resume", return_value="optimized") as optimize_resume,
        ):
            response = self.client.post(
                "/api/resume/optimize",
                json={
                    "resume_text": "  backend developer  ",
                    "target_position": "  Python engineer  ",
                    "job_description": "  FastAPI and Redis required  ",
                },
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["suggestion"], "optimized")
        optimize_resume.assert_called_once_with(
            "backend developer",
            "Python engineer",
            "FastAPI and Redis required",
        )

    def test_resume_generation_uses_job_description(self):
        self.authenticate()
        with (
            patch.object(main, "enforce_llm_limits", return_value=None),
            patch.object(main, "generate_optimized_resume", return_value="generated") as generate_resume,
        ):
            response = self.client.post(
                "/api/resume/generate",
                json={
                    "resume_text": "backend developer",
                    "target_position": "Python engineer",
                    "job_description": "FastAPI and MySQL required",
                },
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["resume"], "generated")
        generate_resume.assert_called_once_with(
            "backend developer",
            "Python engineer",
            "FastAPI and MySQL required",
        )

    def test_interview_and_chat_sessions_are_created_for_user(self):
        user = self.authenticate()
        with patch.object(main, "_save_interview") as save_interview:
            interview_response = self.client.post(
                "/api/interview/start",
                json={
                    "resume_text": "resume",
                    "target_position": "Python engineer",
                    "difficulty": "medium",
                    "job_description": "Build APIs with FastAPI",
                },
            )

        self.assertEqual(interview_response.status_code, 200)
        interview_id = interview_response.json()["data"]["session_id"]
        self.assertTrue(interview_id.startswith("interview_"))
        self.assertEqual(save_interview.call_args.args[2], user.id)
        self.assertEqual(save_interview.call_args.args[1].job_description, "Build APIs with FastAPI")

        with patch.object(main, "_save_chat") as save_chat:
            chat_response = self.client.post("/api/chat/start")

        self.assertEqual(chat_response.status_code, 200)
        chat_id = chat_response.json()["data"]["session_id"]
        self.assertTrue(chat_id.startswith("chat_"))
        self.assertEqual(save_chat.call_args.args[2], user.id)

    def test_interview_report_is_generated_and_saved_for_user(self):
        user = self.authenticate()
        report = {"overall_score": 82, "answered_questions": 2}
        session = SimpleNamespace(report=None, generate_report=Mock(return_value=report))
        session_lock = object()

        with (
            patch.object(main.session_store, "acquire_lock", return_value=session_lock),
            patch.object(main.session_store, "release_lock") as release_lock,
            patch.object(main, "_load_interview", return_value=session),
            patch.object(main, "_save_interview") as save_interview,
            patch.object(main, "save_interview_history") as save_history,
            patch.object(main, "enforce_llm_limits") as enforce_limits,
        ):
            response = self.client.post(
                "/api/interview/report",
                json={"session_id": "interview_demo"},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["report"], report)
        enforce_limits.assert_called_once_with(user)
        save_interview.assert_called_once_with("interview_demo", session, user.id)
        save_history.assert_called_once_with(user.id, "interview_demo", session, report)
        release_lock.assert_called_once_with(session_lock)

    def test_cached_interview_report_does_not_consume_quota(self):
        self.authenticate()
        report = {"overall_score": 90, "answered_questions": 3}
        session = SimpleNamespace(report=report, generate_report=Mock(return_value=report))

        with (
            patch.object(main.session_store, "acquire_lock", return_value=object()),
            patch.object(main.session_store, "release_lock"),
            patch.object(main, "_load_interview", return_value=session),
            patch.object(main, "_save_interview"),
            patch.object(main, "save_interview_history"),
            patch.object(main, "enforce_llm_limits") as enforce_limits,
        ):
            response = self.client.post(
                "/api/interview/report",
                json={"session_id": "interview_cached"},
            )

        self.assertEqual(response.status_code, 200)
        enforce_limits.assert_not_called()

    def test_interview_history_list_is_isolated_to_user(self):
        user = self.authenticate()
        history = SimpleNamespace(
            id="history-1",
            session_id="interview-1",
            target_position="Python工程师",
            difficulty="中级",
            answered_questions=3,
            overall_score=88,
            created_at=None,
        )
        with patch.object(main, "list_interview_histories", return_value=[history]) as list_histories:
            response = self.client.get("/api/interview/history")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["items"][0]["overall_score"], 88)
        list_histories.assert_called_once_with(user.id)

    def test_interview_history_detail_and_delete_remove_redis_session(self):
        user = self.authenticate()
        history = SimpleNamespace(
            id="history-1",
            session_id="interview-1",
            target_position="Python工程师",
            difficulty="中级",
            answered_questions=1,
            overall_score=80,
            created_at=None,
            report_json='{"overall_score": 80, "answered_questions": 1}',
        )
        with patch.object(main, "get_interview_history", return_value=history):
            detail_response = self.client.get("/api/interview/history/history-1")

        self.assertEqual(detail_response.status_code, 200)
        self.assertEqual(detail_response.json()["data"]["report"]["overall_score"], 80)

        with (
            patch.object(main, "delete_interview_history", return_value="interview-1"),
            patch.object(main.session_store, "delete") as delete_session,
        ):
            delete_response = self.client.delete("/api/interview/history/history-1")

        self.assertEqual(delete_response.status_code, 200)
        delete_session.assert_called_once_with("interview", "interview-1")

    def test_delete_all_interview_history_removes_all_redis_sessions(self):
        self.authenticate()
        with (
            patch.object(main, "delete_all_interview_histories", return_value=["interview-1", "interview-2"]),
            patch.object(main.session_store, "delete") as delete_session,
        ):
            response = self.client.delete("/api/interview/history")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["deleted_count"], 2)
        self.assertEqual(delete_session.call_count, 2)

    def test_interview_report_rejects_session_from_another_user(self):
        self.authenticate()
        with (
            patch.object(main.session_store, "acquire_lock", return_value=object()),
            patch.object(main.session_store, "release_lock"),
            patch.object(main, "_load_interview", return_value=None),
        ):
            response = self.client.post(
                "/api/interview/report",
                json={"session_id": "interview_other_user"},
            )

        self.assertEqual(response.status_code, 404)

    def test_knowledge_upload_preserves_user_metadata(self):
        user = self.authenticate()
        vector_module = types.ModuleType("utils.vector_util")
        vector_module.KNOWLEDGE_COLLECTION_NAME = "test_collection"
        vector_module.add_documents = Mock(return_value=2)
        document = SimpleNamespace(id="document-id")

        with (
            patch.dict(sys.modules, {"utils.vector_util": vector_module}),
            patch.object(main, "create_knowledge_document", return_value=document),
            patch.object(main, "update_knowledge_document_chunk_count") as update_count,
        ):
            response = self.client.post(
                "/api/knowledge/upload",
                files={"file": ("questions.txt", "question one\nquestion two", "text/plain")},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["data"]["chunk_count"], 2)
        vector_module.add_documents.assert_called_once_with(
            "test_collection",
            ["question one\nquestion two"],
            metadata={"user_id": str(user.id), "document_id": "document-id"},
        )
        update_count.assert_called_once_with("document-id", user.id, 2)


    def test_knowledge_upload_accepts_pasted_job_duties(self):
        user = self.authenticate()
        vector_module = types.ModuleType("utils.vector_util")
        vector_module.KNOWLEDGE_COLLECTION_NAME = "test_collection"
        vector_module.add_documents = Mock(return_value=1)
        document = SimpleNamespace(id="pasted-document-id")

        with (
            patch.dict(sys.modules, {"utils.vector_util": vector_module}),
            patch.object(main, "create_knowledge_document", return_value=document) as create_document,
            patch.object(main, "update_knowledge_document_chunk_count") as update_count,
        ):
            response = self.client.post(
                "/api/knowledge/upload",
                data={"content": "负责 Python 后端开发，熟悉 FastAPI、MySQL 和 Redis。"},
            )

        self.assertEqual(response.status_code, 200)
        create_document.assert_called_once_with(user.id, "pasted-job-duties.txt")
        vector_module.add_documents.assert_called_once_with(
            "test_collection",
            ["负责 Python 后端开发，熟悉 FastAPI、MySQL 和 Redis。"],
            metadata={"user_id": str(user.id), "document_id": "pasted-document-id"},
        )
        update_count.assert_called_once_with("pasted-document-id", user.id, 1)

    def test_knowledge_upload_accepts_docx_job_duties(self):
        from docx import Document

        user = self.authenticate()
        vector_module = types.ModuleType("utils.vector_util")
        vector_module.KNOWLEDGE_COLLECTION_NAME = "test_collection"
        vector_module.add_documents = Mock(return_value=1)
        document = SimpleNamespace(id="docx-document-id")
        source = BytesIO()
        source_document = Document()
        source_document.add_paragraph("Develop backend APIs with Python and FastAPI.")
        source_document.save(source)

        with (
            patch.dict(sys.modules, {"utils.vector_util": vector_module}),
            patch.object(main, "create_knowledge_document", return_value=document),
            patch.object(main, "update_knowledge_document_chunk_count"),
        ):
            response = self.client.post(
                "/api/knowledge/upload",
                files={
                    "file": (
                        "job-duties.docx",
                        source.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        self.assertEqual(response.status_code, 200)
        vector_module.add_documents.assert_called_once_with(
            "test_collection",
            ["Develop backend APIs with Python and FastAPI."],
            metadata={"user_id": str(user.id), "document_id": "docx-document-id"},
        )

    def test_knowledge_upload_accepts_pdf_job_duties(self):
        from reportlab.pdfgen import canvas

        user = self.authenticate()
        vector_module = types.ModuleType("utils.vector_util")
        vector_module.KNOWLEDGE_COLLECTION_NAME = "test_collection"
        vector_module.add_documents = Mock(return_value=1)
        document = SimpleNamespace(id="pdf-document-id")
        source = BytesIO()
        pdf = canvas.Canvas(source)
        pdf.drawString(72, 720, "Analyze requirements and design backend services.")
        pdf.save()

        with (
            patch.dict(sys.modules, {"utils.vector_util": vector_module}),
            patch.object(main, "create_knowledge_document", return_value=document),
            patch.object(main, "update_knowledge_document_chunk_count"),
        ):
            response = self.client.post(
                "/api/knowledge/upload",
                files={"file": ("job-duties.pdf", source.getvalue(), "application/pdf")},
            )

        self.assertEqual(response.status_code, 200)
        stored_content = vector_module.add_documents.call_args.args[1][0]
        self.assertIn("Analyze requirements", stored_content)
        self.assertEqual(
            vector_module.add_documents.call_args.kwargs["metadata"],
            {"user_id": str(user.id), "document_id": "pdf-document-id"},
        )

    def test_knowledge_upload_requires_text_or_file(self):
        self.authenticate()
        vector_module = types.ModuleType("utils.vector_util")
        vector_module.KNOWLEDGE_COLLECTION_NAME = "test_collection"
        vector_module.add_documents = Mock()

        with patch.dict(sys.modules, {"utils.vector_util": vector_module}):
            response = self.client.post("/api/knowledge/upload", data={"content": ""})

        self.assertEqual(response.status_code, 400)
        vector_module.add_documents.assert_not_called()


if __name__ == "__main__":
    unittest.main()
