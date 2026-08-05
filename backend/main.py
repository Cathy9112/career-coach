import asyncio
import json
import logging
import os
import unicodedata
from io import BytesIO
from pathlib import Path
from urllib.parse import quote
from urllib.request import urlopen
from uuid import uuid4

from fastapi import Depends, FastAPI, HTTPException, File, Form, Query, Request, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from sqlalchemy import text

from services.resume_service import optimize_resume, generate_optimized_resume
from services.interview_service import InterviewSession
from services.chat_service import ChatSession
from auth import (
    User,
    engine,
    authenticate,
    consume_daily_llm_quota,
    create_access_token,
    create_knowledge_document,
    create_user,
    cleanup_expired_daily_usage,
    delete_knowledge_documents,
    delete_knowledge_document,
    get_expired_knowledge_document_ids,
    get_daily_llm_request_count,
    get_current_user,
    delete_all_interview_histories,
    delete_interview_history,
    get_interview_history,
    list_interview_histories,
    save_interview_history,
    update_knowledge_document_chunk_count,
)
from session_store import session_store
from rate_limit import client_ip, rate_limiter
from config import settings


logger = logging.getLogger(__name__)
BASE_DIR = Path(__file__).resolve().parent
TEXT_ENCODINGS = ("utf-8", "gbk", "gb2312", "utf-16")


def _env_int(name: str, default: int) -> int:
    try:
        value = int(os.getenv(name, str(default)))
    except ValueError:
        return default
    return value if value > 0 else default


MAX_UPLOAD_BYTES = _env_int("MAX_UPLOAD_BYTES", 5 * 1024 * 1024)


def _read_upload(file: UploadFile) -> bytes:
    """Read an upload once and reject oversized payloads before parsing."""
    raw_bytes = file.file.read(MAX_UPLOAD_BYTES + 1)
    if len(raw_bytes) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"文件大小不能超过 {MAX_UPLOAD_BYTES} 字节",
        )
    return raw_bytes


def _decode_text(raw_bytes: bytes) -> str:
    for encoding in TEXT_ENCODINGS:
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="文件编码不支持，请另存为 UTF-8 格式后重试")


def parse_text_document(file: UploadFile) -> str:
    """Extract text from a supported resume upload."""
    filename = (file.filename or "").lower()
    if not filename:
        raise HTTPException(status_code=400, detail="缺少文件名")

    if filename.endswith(".txt"):
        content = _decode_text(_read_upload(file))
    elif filename.endswith(".docx"):
        from docx import Document

        doc = Document(BytesIO(_read_upload(file)))
        content = "\n".join(
            paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
        )
    elif filename.endswith(".pdf"):
        import PyPDF2

        reader = PyPDF2.PdfReader(BytesIO(_read_upload(file)))
        page_texts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                page_texts.append(page_text)
        content = "\n".join(page_texts)
    else:
        raise HTTPException(status_code=400, detail="仅支持 .txt、.docx、.pdf 格式文件")

    if not content.strip():
        raise HTTPException(status_code=400, detail="文件内容为空")
    return content


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _document_paragraphs(document) -> list:
    from docx.text.paragraph import Paragraph

    return [
        Paragraph(element, document)
        for element in document.element.body.iterdescendants()
        if element.tag.endswith("}p")
    ]


def _export_docx(original_bytes: bytes, optimized_text: str) -> bytes:
    from docx import Document

    document = Document(BytesIO(original_bytes))
    lines = optimized_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    existing_paragraphs = _document_paragraphs(document)

    for index, paragraph in enumerate(existing_paragraphs):
        _set_paragraph_text(paragraph, lines[index] if index < len(lines) else "")

    for line in lines[len(existing_paragraphs):]:
        document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _wrap_pdf_line(value: str, max_width: int) -> list[str]:
    if not value:
        return [""]
    wrapped: list[str] = []
    current = ""
    current_width = 0
    for char in value:
        char_width = 2 if unicodedata.east_asian_width(char) in {"W", "F", "A"} else 1
        if current and current_width + char_width > max_width:
            wrapped.append(current)
            current = char
            current_width = char_width
        else:
            current += char
            current_width += char_width
    wrapped.append(current)
    return wrapped


def _build_pdf(optimized_text: str, page_width: float, page_height: float) -> bytes:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    margin = 48.0
    font_size = 10.5
    line_height = 17.0
    max_width = max(20, int((page_width - margin * 2) / (font_size * 0.55)))
    embedded_font_name = "CareerCoachChinese"
    configured_font = os.getenv("RESUME_PDF_FONT_PATH", "").strip()
    font_candidates = [
        configured_font,
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "C:/Windows/Fonts/NotoSansSC-VF.ttf",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    font_name = "STSong-Light"
    for font_path in font_candidates:
        if font_path and Path(font_path).is_file():
            if embedded_font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(
                    TTFont(embedded_font_name, font_path, subfontIndex=0)
                )
            font_name = embedded_font_name
            break
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    lines: list[str] = []
    for source_line in optimized_text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        lines.extend(_wrap_pdf_line(source_line, max_width))

    lines_per_page = max(1, int((page_height - margin * 2) / line_height))
    page_chunks = [
        lines[index:index + lines_per_page]
        for index in range(0, len(lines), lines_per_page)
    ] or [[""]]
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=(page_width, page_height), pageCompression=1)
    pdf.setTitle("优化后简历")
    for page_lines in page_chunks:
        text_object = pdf.beginText(margin, page_height - margin)
        text_object.setFont(font_name, font_size)
        text_object.setLeading(line_height)
        for line in page_lines:
            text_object.textLine(line)
        pdf.drawText(text_object)
        pdf.showPage()
    pdf.save()
    return output.getvalue()


def _export_pdf(original_bytes: bytes, optimized_text: str) -> bytes:
    import PyPDF2

    reader = PyPDF2.PdfReader(BytesIO(original_bytes))
    if reader.pages:
        media_box = reader.pages[0].mediabox
        page_width = float(media_box.width)
        page_height = float(media_box.height)
    else:
        page_width, page_height = 595.28, 841.89
    return _build_pdf(optimized_text, page_width, page_height)


def _download_headers(filename: str) -> dict[str, str]:
    ascii_filename = f"optimized_resume{Path(filename).suffix.lower()}"
    encoded_filename = quote(filename)
    return {
        "Content-Disposition": (
            f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{encoded_filename}'
        )
    }


app = FastAPI(title="Career Coach API", version="1.1")
app.mount("/images", StaticFiles(directory=str(BASE_DIR / "images")), name="images")
app.mount("/assets", StaticFiles(directory=str(BASE_DIR / "assets")), name="assets")


def _readiness_status() -> tuple[bool, dict[str, str]]:
    checks: dict[str, str] = {}

    try:
        with engine.connect() as connection:
            connection.execute(text("SELECT 1"))
        checks["database"] = "ok"
    except Exception as exc:
        checks["database"] = f"error:{type(exc).__name__}"

    try:
        session_store.client.ping()
        checks["redis"] = "ok"
    except Exception as exc:
        checks["redis"] = f"error:{type(exc).__name__}"

    chroma_host = os.getenv("CHROMA_HOST")
    if chroma_host:
        chroma_port = int(os.getenv("CHROMA_PORT", "8000"))
        try:
            with urlopen(
                f"http://{chroma_host}:{chroma_port}/api/v2/heartbeat",
                timeout=3,
            ) as response:
                if response.status >= 400:
                    raise RuntimeError("Chroma heartbeat returned an error")
            checks["chroma"] = "ok"
        except Exception as exc:
            checks["chroma"] = f"error:{type(exc).__name__}"
    else:
        checks["chroma"] = "local"

    return all(value in {"ok", "local"} for value in checks.values()), checks


@app.get("/health/live", include_in_schema=False)
async def health_live():
    return {"status": "ok"}


@app.get("/health/ready", include_in_schema=False)
async def health_ready():
    ready, checks = await asyncio.to_thread(_readiness_status)
    status_code = 200 if ready else 503
    return JSONResponse(
        status_code=status_code,
        content={"status": "ready" if ready else "not_ready", "checks": checks},
    )


def cleanup_retained_data() -> None:
    cleanup_expired_daily_usage(settings.usage_retention_days)
    document_ids = get_expired_knowledge_document_ids(settings.knowledge_retention_days)
    if not document_ids:
        return
    try:
        from utils.vector_util import KNOWLEDGE_COLLECTION_NAME, delete_documents

        delete_documents(KNOWLEDGE_COLLECTION_NAME, document_ids)
    except Exception:
        logger.exception("Knowledge retention cleanup failed")
        return
    delete_knowledge_documents(document_ids)


async def run_retention_cleanup() -> None:
    while True:
        await asyncio.sleep(settings.retention_cleanup_interval_seconds)
        await asyncio.to_thread(cleanup_retained_data)


@app.on_event("startup")
async def start_retention_cleanup() -> None:
    await asyncio.to_thread(cleanup_retained_data)
    app.state.retention_cleanup_task = asyncio.create_task(run_retention_cleanup())


@app.on_event("shutdown")
async def stop_retention_cleanup() -> None:
    task = getattr(app.state, "retention_cleanup_task", None)
    if task:
        task.cancel()


@app.get("/")
async def serve_index():
    """
    项目根路由访问接口
    访问 http://127.0.0.1:5200 直接返回前端首页index.html
    实现前后端一体化部署，无需单独启动前端服务
    """
    return FileResponse(str(BASE_DIR / "index.html"))


configured_origins = [
    origin.strip()
    for origin in os.getenv(
        "CORS_ORIGINS", "http://127.0.0.1:5200,http://localhost:5200"
    ).split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=configured_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ========== 请求参数校验模型 ==========
# Pydantic实体类，接收简历优化接口入参，自动校验字段类型、缺失参数
class ResumeOptimizeReq(BaseModel):
    resume_text: str = Field(..., min_length=1, max_length=50_000)
    target_position: str = Field(
        default="Java后端开发工程师", min_length=1, max_length=200
    )
    job_description: str = Field(default="", max_length=20_000)


# Pydantic实体类，开启模拟面试接口的入参校验模型
class InterviewStartReq(BaseModel):
    resume_text: str = Field(..., min_length=1, max_length=50_000)
    target_position: str = Field(..., min_length=1, max_length=200)
    job_description: str = Field(default="", max_length=20_000)


class InterviewStreamReq(BaseModel):
    session_id: str = Field(..., min_length=1, max_length=100)
    answer: str = Field(default="", max_length=20_000)


class InterviewReportReq(BaseModel):
    session_id: str = Field(..., min_length=1, max_length=100)


class ChatStreamReq(BaseModel):
    session_id: str = Field(..., min_length=1, max_length=100)
    message: str = Field(..., min_length=1, max_length=20_000)


class AuthReq(BaseModel):
    username: str = Field(..., min_length=3, max_length=120)
    password: str = Field(..., min_length=8, max_length=200)


def _clean_text(value: str, field_name: str) -> str:
    cleaned = value.strip()
    if not cleaned:
        raise HTTPException(status_code=422, detail=f"{field_name}不能为空")
    return cleaned


def _clean_optional_text(value: str) -> str:
    return value.strip()


def _sse(payload: dict) -> str:
    return f"data:{json.dumps(payload, ensure_ascii=False)}\n\n"


SSE_HEADERS = {
    "Cache-Control": "no-cache",
    "Connection": "keep-alive",
    "X-Accel-Buffering": "no",
}


def _next_session_id(prefix: str) -> str:
    return f"{prefix}_{uuid4().hex}"


def _save_interview(session_id: str, session: InterviewSession, user_id: int) -> None:
    session_store.save("interview", session_id, {
        "resume_text": session.resume_text,
        "target_position": session.target_position,
        "job_description": session.job_description,
        "chat_list": session.chat_list,
        "qa_history": session.qa_history,
        "report": session.report,
        "user_id": user_id,
    })


def _load_interview(session_id: str, user_id: int) -> InterviewSession | None:
    payload = session_store.load("interview", session_id)
    if not payload or payload.get("user_id") != user_id:
        return None
    session = InterviewSession(
        payload["resume_text"],
        payload["target_position"],
        user_id,
        payload.get("job_description", ""),
    )
    session.chat_list = payload.get("chat_list", session.chat_list)
    session.qa_history = payload.get("qa_history", [])
    session.report = payload.get("report")
    return session


def _save_chat(session_id: str, session: ChatSession, user_id: int) -> None:
    session_store.save("chat", session_id, {"chat_list": session.chat_list, "user_id": user_id})


def _load_chat(session_id: str, user_id: int) -> ChatSession | None:
    payload = session_store.load("chat", session_id)
    if not payload or payload.get("user_id") != user_id:
        return None
    session = ChatSession()
    session.chat_list = payload.get("chat_list", session.chat_list)
    return session


def _auth_response(response: Response, token: str) -> dict:
    response.set_cookie(
        "access_token",
        token,
        httponly=True,
        secure=settings.cookie_secure,
        samesite="lax",
        max_age=8 * 3600,
    )
    return {"code": 200, "data": {"access_token": token, "token_type": "bearer"}}


def enforce_llm_limits(user: User) -> None:
    rate_limiter.enforce(
        "llm",
        str(user.id),
        settings.llm_rate_limit,
        settings.llm_rate_window_seconds,
    )
    if not consume_daily_llm_quota(user.id, settings.daily_llm_request_quota):
        raise HTTPException(status_code=429, detail="今日模型调用额度已用尽")


@app.post("/api/auth/register")
def api_register(req: AuthReq, response: Response, request: Request):
    rate_limiter.enforce("auth", client_ip(request), settings.auth_rate_limit, settings.auth_rate_window_seconds)
    try:
        user = create_user(req.username.strip(), req.password)
    except ValueError as exc:
        raise HTTPException(status_code=409, detail="username already exists") from exc
    return _auth_response(response, create_access_token(user))


@app.post("/api/auth/login")
def api_login(req: AuthReq, response: Response, request: Request):
    rate_limiter.enforce("auth", client_ip(request), settings.auth_rate_limit, settings.auth_rate_window_seconds)
    user = authenticate(req.username.strip(), req.password)
    if not user:
        raise HTTPException(status_code=401, detail="invalid username or password")
    return _auth_response(response, create_access_token(user))


@app.post("/api/auth/logout")
def api_logout(response: Response):
    response.delete_cookie(
        "access_token",
        secure=settings.cookie_secure,
        httponly=True,
        samesite="lax",
    )
    return {"code": 200, "data": {"message": "logged out"}}


@app.get("/api/auth/me")
def api_me(user: User = Depends(get_current_user)):
    return {"code": 200, "data": {"id": user.id, "username": user.username}}


@app.get("/api/usage")
def api_usage(user: User = Depends(get_current_user)):
    used = get_daily_llm_request_count(user.id)
    quota = settings.daily_llm_request_quota
    return {"code": 200, "data": {"daily_llm_quota": quota, "used": used, "remaining": max(quota - used, 0)}}


# ========== 1. 获取简历优化诊断建议接口 ==========
@app.post("/api/resume/optimize")
def api_resume_optimize(req: ResumeOptimizeReq, user: User = Depends(get_current_user)):
    """
    接口功能：根据简历+岗位生成三大模块优化建议
    请求方式：POST 向服务器提交数据
    入参：简历文本、目标岗位
    返回：量化改写、动词替换、结构补充三类优化方案
    """
    try:
        enforce_llm_limits(user)
        result = optimize_resume(
            _clean_text(req.resume_text, "resume_text"),
            _clean_text(req.target_position, "target_position"),
            _clean_optional_text(req.job_description),
        )
        return {"code": 200, "data": {"suggestion": result}}
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Resume optimization failed")
        raise HTTPException(status_code=500, detail="简历优化服务暂时不可用") from exc


# ========== 2. 一键生成完整优化简历接口 ==========
@app.post("/api/resume/generate")
def api_resume_generate(req: ResumeOptimizeReq, user: User = Depends(get_current_user)):
    """
    接口功能：基于原始简历，生成适配岗位的完整标准化简历
    请求方式：POST
    """
    try:
        enforce_llm_limits(user)
        result = generate_optimized_resume(
            _clean_text(req.resume_text, "resume_text"),
            _clean_text(req.target_position, "target_position"),
            _clean_optional_text(req.job_description),
        )
        return {"code": 200, "data": {"resume": result}}
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Resume generation failed")
        raise HTTPException(status_code=500, detail="简历生成服务暂时不可用") from exc


# ========== 3. 创建模拟面试会话接口 ==========
@app.post("/api/interview/start")
def api_interview_start(req: InterviewStartReq, user: User = Depends(get_current_user)):
    """
    接口功能：初始化一场独立模拟面试，生成唯一会话ID
    请求方式：POST
    入参：简历、岗位和可选岗位JD
    返回：session_id，后续流式对话依靠该ID区分用户会话
    """
    resume_text = _clean_text(req.resume_text, "resume_text")
    target_position = _clean_text(req.target_position, "target_position")
    job_description = _clean_optional_text(req.job_description)
    session_id = _next_session_id("interview")
    session = InterviewSession(
        resume_text=resume_text,
        target_position=target_position,
        user_id=user.id,
        job_description=job_description,
    )
    _save_interview(session_id, session, user.id)
    return {"code": 200, "data": {"session_id": session_id}}


# ========== 4. 模拟面试SSE流式问答接口 ==========
@app.post("/api/interview/stream")
def api_interview_stream(
    req: InterviewStreamReq,
    user: User = Depends(get_current_user),
):
    """
    接口功能：长连接流式返回面试官逐字提问，实现打字机效果
    请求方式：POST，敏感回答通过 JSON 请求体传输
    入参：会话ID、用户上一轮回答
    """
    session_id = req.session_id
    answer = req.answer
    # 校验会话是否存在，不存在抛出404
    session_lock = session_store.acquire_lock("interview", session_id)
    if session_lock is None:
        raise HTTPException(status_code=409, detail="会话正在处理中，请勿重复提交")

    try:
        session = _load_interview(session_id, user.id)
    except Exception:
        session_store.release_lock(session_lock)
        raise
    if session is None:
        session_store.release_lock(session_lock)
        raise HTTPException(status_code=404, detail="会话不存在")
    try:
        enforce_llm_limits(user)
    except Exception:
        session_store.release_lock(session_lock)
        raise

    # 生成器函数：循环逐段产出大模型返回文本
    def generate():
        try:
            # 迭代流式文本分片
            for text in session.stream_next_question(answer):
                yield _sse({"content": text})
            _save_interview(session_id, session, user.id)
            yield "data:[DONE]\n\n"
        except Exception:
            logger.exception("Interview stream failed")
            yield _sse({"error": "面试服务暂时不可用，请稍后重试"})
            yield "data:[DONE]\n\n"
        finally:
            session_store.release_lock(session_lock)

    # 返回SSE长连接流式响应，设置长连接专用请求头
    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers=SSE_HEADERS,
    )


# ========== 5. 生成面试评分和最终报告接口 ==========
@app.post("/api/interview/report")
def api_interview_report(
    req: InterviewReportReq,
    user: User = Depends(get_current_user),
):
    session_lock = session_store.acquire_lock("interview", req.session_id)
    if session_lock is None:
        raise HTTPException(status_code=409, detail="会话正在处理中，请稍后生成报告")

    try:
        session = _load_interview(req.session_id, user.id)
        if session is None:
            raise HTTPException(status_code=404, detail="会话不存在")
        if session.report is None:
            enforce_llm_limits(user)
        try:
            report = session.generate_report()
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        _save_interview(req.session_id, session, user.id)
        save_interview_history(user.id, req.session_id, session, report)
        return {"code": 200, "data": {"report": report}}
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Interview report generation failed")
        raise HTTPException(status_code=500, detail="面试报告生成失败，请稍后重试") from exc
    finally:
        session_store.release_lock(session_lock)


# ========== 6. 面试历史记录与数据删除接口 ==========
def _history_summary(history) -> dict:
    return {
        "id": history.id,
        "session_id": history.session_id,
        "target_position": history.target_position,
        "answered_questions": history.answered_questions,
        "overall_score": history.overall_score,
        "created_at": history.created_at.isoformat() if history.created_at else None,
    }


@app.get("/api/interview/history")
def api_interview_history(user: User = Depends(get_current_user)):
    histories = list_interview_histories(user.id)
    return {"code": 200, "data": {"items": [_history_summary(item) for item in histories]}}


@app.get("/api/interview/history/{history_id}")
def api_interview_history_detail(history_id: str, user: User = Depends(get_current_user)):
    import json

    history = get_interview_history(history_id, user.id)
    if history is None:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    try:
        report = json.loads(history.report_json)
    except (TypeError, json.JSONDecodeError) as exc:
        logger.exception("Invalid interview history report", exc_info=exc)
        raise HTTPException(status_code=500, detail="历史报告数据损坏") from exc
    return {"code": 200, "data": {"history": _history_summary(history), "report": report}}


@app.delete("/api/interview/history/{history_id}")
def api_delete_interview_history(history_id: str, user: User = Depends(get_current_user)):
    session_id = delete_interview_history(history_id, user.id)
    if session_id is None:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    session_store.delete("interview", session_id)
    return {"code": 200, "data": {"message": "历史记录已删除"}}


@app.delete("/api/interview/history")
def api_delete_all_interview_history(user: User = Depends(get_current_user)):
    session_ids = delete_all_interview_histories(user.id)
    for session_id in session_ids:
        session_store.delete("interview", session_id)
    return {"code": 200, "data": {"deleted_count": len(session_ids)}}


# ========== 7. 创建通用AI聊天会话接口 ==========
@app.post("/api/chat/start")
def api_chat_start(user: User = Depends(get_current_user)):
    """
    接口功能：创建通用求职问答会话，生成独立session_id
    请求方式：POST
    """
    session_id = _next_session_id("chat")
    session = ChatSession()
    _save_chat(session_id, session, user.id)
    return {"code": 200, "data": {"session_id": session_id}}


# ========== 8. 通用AI聊天SSE流式问答接口 ==========
@app.post("/api/chat/stream")
def api_chat_stream(
    req: ChatStreamReq,
    user: User = Depends(get_current_user),
):
    """
    接口功能：通用职业助手流式对话，逻辑与面试流式接口完全复用
    请求方式：POST，敏感消息通过 JSON 请求体传输
    """
    session_id = req.session_id
    message = req.message
    message = _clean_text(message, "message")
    session_lock = session_store.acquire_lock("chat", session_id)
    if session_lock is None:
        raise HTTPException(status_code=409, detail="会话正在处理中，请勿重复提交")

    try:
        session = _load_chat(session_id, user.id)
    except Exception:
        session_store.release_lock(session_lock)
        raise
    if session is None:
        session_store.release_lock(session_lock)
        raise HTTPException(status_code=404, detail="会话不存在")
    try:
        enforce_llm_limits(user)
    except Exception:
        session_store.release_lock(session_lock)
        raise

    # 流式文本生成器
    def generate():
        try:
            for text in session.stream_reply(message):
                yield _sse({"content": text})
            _save_chat(session_id, session, user.id)
            yield "data:[DONE]\n\n"
        except Exception:
            logger.exception("Chat stream failed")
            yield _sse({"error": "聊天服务暂时不可用，请稍后重试"})
            yield "data:[DONE]\n\n"
        finally:
            session_store.release_lock(session_lock)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers=SSE_HEADERS,
    )


# ========== 7. 简历文件上传解析接口 ==========
@app.post("/api/resume/upload")
def api_resume_upload(file: UploadFile = File(...), _: User = Depends(get_current_user)):
    """
    接口功能：接收用户上传简历文件，调用通用解析函数提取纯文本
    请求方式：POST 表单FormData文件上传,提交二进制文件
    """
    try:
        content = parse_text_document(file)
        return {"code": 200, "data": {"content": content}}
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Resume upload failed")
        raise HTTPException(status_code=400, detail="无法解析简历文件") from exc


# ========== 8. 优化简历导出接口 ==========
@app.post("/api/resume/export")
def api_resume_export(
    optimized_text: str = Form(..., min_length=1, max_length=50_000),
    file: UploadFile | None = File(None),
    _: User = Depends(get_current_user),
):
    cleaned_text = _clean_text(optimized_text, "optimized_text")
    if file is None or not file.filename:
        filename = "优化后简历.txt"
        return Response(
            content=cleaned_text.encode("utf-8-sig"),
            media_type="text/plain; charset=utf-8",
            headers=_download_headers(filename),
        )

    source_name = file.filename.replace("\\", "/").split("/")[-1]
    extension = Path(source_name).suffix.lower()
    if extension not in {".txt", ".docx", ".pdf"}:
        raise HTTPException(status_code=400, detail="仅支持导出 .txt、.docx、.pdf 格式文件")

    original_bytes = _read_upload(file)
    output_name = f"{Path(source_name).stem}_优化版{extension}"
    try:
        if extension == ".txt":
            exported_bytes = cleaned_text.encode("utf-8-sig")
            media_type = "text/plain; charset=utf-8"
        elif extension == ".docx":
            exported_bytes = _export_docx(original_bytes, cleaned_text)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            exported_bytes = _export_pdf(original_bytes, cleaned_text)
            media_type = "application/pdf"
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Resume export failed")
        raise HTTPException(status_code=400, detail="无法按源文件格式导出简历") from exc

    return Response(
        content=exported_bytes,
        media_type=media_type,
        headers=_download_headers(output_name),
    )


# ========== 9. 面试知识库文档上传入库接口 ==========
@app.post("/api/knowledge/upload")
def api_knowledge_upload(
    file: UploadFile | None = File(None),
    content: str = Form(""),
    user: User = Depends(get_current_user),
):
    """Store pasted or uploaded job responsibilities in the current user's vector collection."""
    document_id = None
    try:
        from utils.vector_util import KNOWLEDGE_COLLECTION_NAME, add_documents

        content_parts: list[str] = []
        filename = ""
        if file is not None and file.filename:
            filename = file.filename
            content_parts.append(parse_text_document(file).strip())

        pasted_content = content.strip()
        if pasted_content:
            if len(pasted_content.encode("utf-8")) > MAX_UPLOAD_BYTES:
                raise HTTPException(status_code=413, detail="粘贴内容大小超过限制")
            content_parts.append(pasted_content)

        if not content_parts:
            raise HTTPException(status_code=400, detail="请粘贴岗位职责或上传文件")

        combined_content = "\n\n".join(part for part in content_parts if part)
        if not combined_content:
            raise HTTPException(status_code=400, detail="岗位职责内容不能为空")

        source_name = filename or "pasted-job-duties.txt"
        document = create_knowledge_document(user.id, source_name)
        document_id = document.id
        chunk_count = add_documents(
            KNOWLEDGE_COLLECTION_NAME,
            [combined_content],
            metadata={"user_id": str(user.id), "document_id": document.id},
        )
        update_knowledge_document_chunk_count(document.id, user.id, chunk_count)

        return {
            "code": 200,
            "data": {
                "filename": source_name,
                "document_id": document.id,
                "chunk_count": chunk_count,
                "message": f"岗位职责保存成功，共生成 {chunk_count} 个内容片段。模拟面试将结合这些内容和您的简历进行提问。",
            },
        }
    except HTTPException:
        raise
    except Exception as exc:
        if document_id:
            delete_knowledge_document(document_id, user.id)
        logger.exception("Job duties upload failed")
        raise HTTPException(status_code=503, detail="岗位职责保存失败，请稍后重试") from exc


@app.get("/api/knowledge/query")
def api_knowledge_query(
    query: str = Query(..., min_length=1, max_length=500),
    n_results: int = Query(3, ge=1, le=20),
    user: User = Depends(get_current_user),
):
    """
    接口功能：输入关键词，测试向量库相似度检索效果
    请求方式：GET
    参数：query检索关键词，n_results返回匹配知识点条数（默认3条）
    """
    try:
        # 懒导入向量检索工具
        from utils.vector_util import KNOWLEDGE_COLLECTION_NAME, query_documents
        # 在面试题库集合中检索相似文本片段
        results = query_documents(
            KNOWLEDGE_COLLECTION_NAME,
            query.strip(),
            n_results,
            where={"user_id": str(user.id)},
        )
        return {"code": 200, "data": {"results": results}}
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Knowledge base query failed")
        raise HTTPException(status_code=503, detail="知识库服务暂时不可用") from exc


if __name__ == "__main__":
    raise SystemExit(
        "Do not run main.py directly. Use uvicorn main:app --reload for development "
        "or Docker Compose for production."
    )
